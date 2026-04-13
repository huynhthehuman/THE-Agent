import os
from docx import Document
from docx.shared import Pt
import re

def markdown_to_docx(md_path, docx_path):
    doc = Document()
    # Thêm tiêu đề mặc định
    title = doc.add_heading('THÔNG BÁO: CHUẨN HOÁ TRÁCH NHIỆM & LUỒNG PHỐI HỢP LIÊN BAN', level=0)
    
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    skip_next = False
    for line in lines:
        line = line.strip('\n')
        if not line or line == '---':
            continue
            
        if line.startswith('# THÔNG BÁO'):
            continue # Đã add header trên đầu
            
        if line.startswith('## '):
            doc.add_heading(line[3:], level=1)
        elif line.startswith('### '):
            doc.add_heading(line[4:], level=2)
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            text = line[2:]
            render_inline(p, text)
        else:
            p = doc.add_paragraph()
            render_inline(p, line)

    doc.save(docx_path)
    print("Exported to:", docx_path)

def render_inline(paragraph, text):
    # Regex tách in đậm
    chunks = re.split(r'(\*\*.*?\*\*)', text)
    for chunk in chunks:
        if chunk.startswith('**') and chunk.endswith('**'):
            run = paragraph.add_run(chunk[2:-2])
            run.bold = True
        else:
            paragraph.add_run(chunk)

if __name__ == '__main__':
    md_file = r'C:\Users\leovn\.gemini\antigravity\brain\274f66bd-a0bf-4b6c-b554-125491f12d25\phan_quyen_he_thong.md'
    docx_file = r'C:\Users\leovn\Documents\THE_Agent_Macbook_Export\Thong_Bao_Phan_Quyen.docx'
    markdown_to_docx(md_file, docx_file)
